"""
MKA Template Renderer — PDF/Word 匯出。

對照 ENGINEERING_PLAN.md §7.1、§5.5：
- Fixed Form 不只是自由生成 Markdown
- export 需符合 approval policy
- 正式文件 watermark/version
- PDF/Word/Excel 版型

首版實作：
- Markdown → PDF（用 reportlab）
- Markdown → Word（用 python-docx）
- 水印（版本、核准者、時間戳）
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


@dataclass
class ExportResult:
    """匯出結果。"""
    format: str = ""  # pdf | docx | xlsx | md
    content: bytes = b""
    filename: str = ""
    error: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def success(self) -> bool:
        return not self.error and bool(self.content)


class TemplateRenderer:
    """表單匯出引擎。"""

    def render_pdf(
        self,
        title: str,
        fields: Dict[str, Any],
        provenance: Dict[str, Any],
        approval_info: Dict[str, Any],
    ) -> ExportResult:
        """渲染 PDF。

        注意：reportlab 預設 Helvetica 不支援 CJK。
        生產環境需註冊 CJK 字型（如 STSong-Light 或 Noto Sans CJK）。
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            )
            from reportlab.lib import colors

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm)
            styles = getSampleStyleSheet()
            story: List[Any] = []

            # 嘗試註冊 CJK 字型
            cjk_font_name = self._register_cjk_font()
            if cjk_font_name:
                # 建立新的 CJK 樣式，不修改原有 styles（避免破壞 reportlab 內部）
                cjk_title = ParagraphStyle(
                    "CJKTitle", parent=styles["Title"], fontName=cjk_font_name
                )
                cjk_normal = ParagraphStyle(
                    "CJKNormal", parent=styles["Normal"], fontName=cjk_font_name
                )
                cjk_heading = ParagraphStyle(
                    "CJKHeading", parent=styles["Heading2"], fontName=cjk_font_name
                )
            else:
                cjk_title = styles["Title"]
                cjk_normal = styles["Normal"]
                cjk_heading = styles["Heading2"]
                logger.warning("CJK font not available; PDF may not display Chinese correctly")

            # 標題
            story.append(Paragraph(title, cjk_title))
            story.append(Spacer(1, 10*mm))

            # 欄位表
            field_data = [["欄位", "值"]]
            for name, value in fields.items():
                field_data.append([name, str(value)])

            table = Table(field_data, colWidths=[50*mm, 120*mm])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(table)
            story.append(Spacer(1, 10*mm))

            # 核准資訊
            if approval_info:
                story.append(Paragraph("核准資訊", cjk_heading))
                for key, value in approval_info.items():
                    story.append(Paragraph(f"<b>{key}:</b> {value}", cjk_normal))
                story.append(Spacer(1, 5*mm))

            # 水印（版本 + 時間戳）
            timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
            version = approval_info.get("version", "1.0")
            story.append(Spacer(1, 15*mm))
            watermark_style = ParagraphStyle(
                "Watermark",
                parent=cjk_normal,
                fontSize=8,
                textColor=colors.grey,
                alignment=2,  # right
            )
            story.append(Paragraph(
                f"Version {version} | Generated {timestamp} | Enclave MKA",
                watermark_style
            ))

            doc.build(story)
            pdf_bytes = buffer.getvalue()

            return ExportResult(
                format="pdf",
                content=pdf_bytes,
                filename=f"{title}_v{version}.pdf",
                metadata={"pages": 1, "version": version},
            )

        except ImportError:
            return ExportResult(format="pdf", error="reportlab not installed")
        except Exception as exc:
            logger.error(f"PDF render failed: {exc}")
            return ExportResult(format="pdf", error=str(exc))

    def render_docx(
        self,
        title: str,
        fields: Dict[str, Any],
        provenance: Dict[str, Any],
        approval_info: Dict[str, Any],
    ) -> ExportResult:
        """渲染 Word DOCX。"""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()

            # 標題
            doc.add_heading(title, level=0)

            # 欄位表
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "欄位"
            hdr[1].text = "值"

            for name, value in fields.items():
                row = table.add_row().cells
                row[0].text = name
                row[1].text = str(value)

            doc.add_paragraph("")

            # 核准資訊
            if approval_info:
                doc.add_heading("核准資訊", level=2)
                for key, value in approval_info.items():
                    doc.add_paragraph(f"{key}: {value}")

            # 水印
            timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
            version = approval_info.get("version", "1.0")
            p = doc.add_paragraph()
            p.alignment = 2  # right
            run = p.add_run(f"Version {version} | Generated {timestamp} | Enclave MKA")
            run.font.size = Pt(8)
            # 設定灰色字
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(128, 128, 128)

            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()

            return ExportResult(
                format="docx",
                content=docx_bytes,
                filename=f"{title}_v{version}.docx",
                metadata={"version": version},
            )

        except ImportError:
            return ExportResult(format="docx", error="python-docx not installed")
        except Exception as exc:
            logger.error(f"DOCX render failed: {exc}")
            return ExportResult(format="docx", error=str(exc))

    def render_markdown(
        self,
        title: str,
        fields: Dict[str, Any],
        provenance: Dict[str, Any],
        approval_info: Dict[str, Any],
    ) -> ExportResult:
        """渲染 Markdown。"""
        lines = [f"# {title}\n"]

        # 欄位
        lines.append("## 欄位\n")
        lines.append("| 欄位 | 值 |")
        lines.append("|------|-----|")
        for name, value in fields.items():
            lines.append(f"| {name} | {value} |")
        lines.append("")

        # 核准資訊
        if approval_info:
            lines.append("## 核准資訊\n")
            for key, value in approval_info.items():
                lines.append(f"- **{key}**: {value}")
            lines.append("")

        # 水印
        timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        version = approval_info.get("version", "1.0")
        lines.append(f"\n---\n*Version {version} | Generated {timestamp} | Enclave MKA*")

        md_content = "\n".join(lines).encode("utf-8")
        return ExportResult(
            format="md",
            content=md_content,
            filename=f"{title}_v{version}.md",
            metadata={"version": version},
        )

    def render_excel(
        self,
        title: str,
        fields: Dict[str, Any],
        provenance: Dict[str, Any],
        approval_info: Dict[str, Any],
    ) -> ExportResult:
        """渲染 Excel XLSX。"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill

            wb = Workbook()
            ws = wb.active
            ws.title = title[:31]  # Excel sheet name max 31 chars

            # 標題列
            ws["A1"] = title
            ws["A1"].font = Font(size=14, bold=True)
            ws.merge_cells("A1:B1")

            # 欄位表
            ws["A3"] = "欄位"
            ws["B3"] = "值"
            ws["A3"].font = Font(bold=True)
            ws["B3"].font = Font(bold=True)
            ws["A3"].fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            ws["B3"].fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")

            row = 4
            for name, value in fields.items():
                ws[f"A{row}"] = name
                ws[f"B{row}"] = str(value)
                row += 1

            # 核准資訊
            if approval_info:
                row += 1
                ws[f"A{row}"] = "核准資訊"
                ws[f"A{row}"].font = Font(bold=True)
                row += 1
                for key, value in approval_info.items():
                    ws[f"A{row}"] = key
                    ws[f"B{row}"] = str(value)
                    row += 1

            # 水印
            row += 1
            timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
            version = approval_info.get("version", "1.0")
            ws[f"A{row}"] = f"Version {version} | Generated {timestamp} | Enclave MKA"
            ws[f"A{row}"].font = Font(size=8, color="888888")

            # 欄寬
            ws.column_dimensions["A"].width = 25
            ws.column_dimensions["B"].width = 40

            buffer = io.BytesIO()
            wb.save(buffer)
            xlsx_bytes = buffer.getvalue()

            return ExportResult(
                format="xlsx",
                content=xlsx_bytes,
                filename=f"{title}_v{version}.xlsx",
                metadata={"version": version},
            )

        except ImportError:
            return ExportResult(format="xlsx", error="openpyxl not installed")
        except Exception as exc:
            logger.error(f"Excel render failed: {exc}")
            return ExportResult(format="xlsx", error=str(exc))

    def _register_cjk_font(self) -> str:
        """註冊 CJK 字型供 reportlab 使用。

        嘗試順序：
        1. reportlab 內建 CID 字型（STSong-Light）
        2. 系統已安裝的 CJK 字型（Noto Sans CJK、微軟正黑體）

        Returns:
            註冊成功的字型名稱，或空字串（失敗）
        """
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont

            # 嘗試內建 CID 字型
            try:
                pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
                return "STSong-Light"
            except Exception:
                pass

            # 嘗試系統字型
            import os
            font_paths = [
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "C:/Windows/Fonts/msyh.ttc",  # 微軟正黑體
                "C:/Windows/Fonts/simsun.ttc",  # 宋體
            ]
            from reportlab.pdfbase.ttfonts import TTFont
            for path in font_paths:
                if os.path.exists(path):
                    try:
                        pdfmetrics.registerFont(TTFont("CJK", path))
                        return "CJK"
                    except Exception:
                        continue

            return ""
        except ImportError:
            return ""


# ── 單例 ──

_renderer: Optional[TemplateRenderer] = None


def get_template_renderer() -> TemplateRenderer:
    global _renderer
    if _renderer is None:
        _renderer = TemplateRenderer()
    return _renderer