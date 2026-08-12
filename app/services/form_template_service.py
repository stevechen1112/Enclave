"""Company DOCX/XLSX template upload, placeholder parse, mapping, preview."""
from __future__ import annotations

import io
import logging
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID, uuid4

from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_.-]+)\s*\}\}")


def convert_office_template_to_pdf(content: bytes, source_format: str, stem: str) -> bytes:
    """Convert a rendered company DOCX/XLSX template to PDF with LibreOffice."""
    extension = source_format.lower().lstrip(".")
    if extension not in {"docx", "xlsx"}:
        raise ValueError(f"unsupported office template format: {source_format}")
    safe_stem = re.sub(r"[^A-Za-z0-9_.-]", "_", stem or "form")
    with tempfile.TemporaryDirectory(prefix="enclave-form-pdf-") as temp_dir:
        root = Path(temp_dir)
        source = root / f"{safe_stem}.{extension}"
        output_dir = root / "output"
        profile_dir = root / "profile"
        output_dir.mkdir()
        profile_dir.mkdir()
        source.write_bytes(content)
        completed = subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation={profile_dir.as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(source),
            ],
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
        )
        pdf = output_dir / f"{safe_stem}.pdf"
        if completed.returncode != 0 or not pdf.exists() or pdf.stat().st_size == 0:
            detail = (completed.stderr or completed.stdout).strip()[:500]
            raise RuntimeError(f"company template PDF conversion failed: {detail or 'no PDF output'}")
        return pdf.read_bytes()


def extract_placeholders_docx(content: bytes) -> List[str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx required for DOCX templates") from exc
    doc = Document(io.BytesIO(content))
    found: List[str] = []
    texts: List[str] = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    for text in texts:
        for m in _PLACEHOLDER_RE.finditer(text or ""):
            key = m.group(1)
            if key not in found:
                found.append(key)
    return found


def extract_placeholders_xlsx(content: bytes) -> List[str]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl required for XLSX templates") from exc
    wb = load_workbook(io.BytesIO(content))
    found: List[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if not isinstance(cell, str):
                    continue
                for m in _PLACEHOLDER_RE.finditer(cell):
                    key = m.group(1)
                    if key not in found:
                        found.append(key)
    return found


def render_docx_template(content: bytes, values: Dict[str, Any], mapping: Dict[str, str]) -> bytes:
    from docx import Document

    doc = Document(io.BytesIO(content))

    def _replace(text: str) -> str:
        if not text:
            return text
        out = text
        for ph, field_key in mapping.items():
            token = "{{" + ph + "}}"
            # also allow mapped key as placeholder
            alt = "{{" + field_key + "}}"
            val = str(values.get(field_key, values.get(ph, "")))
            out = out.replace(token, val).replace(alt, val)
        # direct {{field}} without explicit mapping
        for m in list(_PLACEHOLDER_RE.finditer(out)):
            key = m.group(1)
            if key in values:
                out = out.replace(m.group(0), str(values[key]))
        return out

    for p in doc.paragraphs:
        if "{{" in (p.text or ""):
            p.text = _replace(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in (cell.text or ""):
                    cell.text = _replace(cell.text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_xlsx_template(content: bytes, values: Dict[str, Any], mapping: Dict[str, str]) -> bytes:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content))

    def _replace(text: str) -> str:
        out = text
        for ph, field_key in mapping.items():
            token = "{{" + ph + "}}"
            alt = "{{" + field_key + "}}"
            val = str(values.get(field_key, values.get(ph, "")))
            out = out.replace(token, val).replace(alt, val)
        for m in list(_PLACEHOLDER_RE.finditer(out)):
            key = m.group(1)
            if key in values:
                out = out.replace(m.group(0), str(values[key]))
        return out

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and "{{" in cell.value:
                    cell.value = _replace(cell.value)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class FormTemplateService:
    def __init__(self, db: Session):
        self.db = db

    def _storage(self):
        from app.services.storage import get_storage_backend
        return get_storage_backend()

    def list_templates(self, tenant_id: UUID, form_key: Optional[str] = None) -> List[Any]:
        from app.models.mka import FormTemplate
        q = self.db.query(FormTemplate).filter(FormTemplate.tenant_id == tenant_id)
        if form_key:
            q = q.filter(FormTemplate.form_key == form_key)
        return q.order_by(FormTemplate.form_key.asc(), FormTemplate.created_at.desc()).all()

    def upload(
        self,
        *,
        tenant_id: UUID,
        user_id: UUID,
        form_key: str,
        name: str,
        filename: str,
        content: bytes,
        version: str = "1.0",
    ) -> Any:
        from app.models.mka import FormTemplate

        lower = filename.lower()
        if lower.endswith(".docx"):
            fmt = "docx"
            placeholders = extract_placeholders_docx(content)
        elif lower.endswith(".xlsx"):
            fmt = "xlsx"
            placeholders = extract_placeholders_xlsx(content)
        else:
            raise ValueError("only .docx or .xlsx supported")

        import os
        import tempfile
        from app.services.storage import build_storage_key

        storage_key = build_storage_key(tenant_id, uuid4(), fmt)
        fd, tmp_path = tempfile.mkstemp(suffix=f".{fmt}")
        try:
            with os.fdopen(fd, "wb") as fh:
                fh.write(content)
            self._storage().put(storage_key, tmp_path)
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        # default identity mapping
        mapping = {ph: ph for ph in placeholders}
        row = FormTemplate(
            tenant_id=tenant_id,
            form_key=form_key,
            name=name or filename,
            format=fmt,
            version=version,
            storage_key=storage_key,
            placeholders=placeholders,
            field_mapping=mapping,
            status="draft",
            created_by=user_id,
        )
        self.db.add(row)
        self.db.flush()
        return row

    def update_mapping(self, *, tenant_id: UUID, template_id: UUID, mapping: Dict[str, str]) -> Any:
        from app.models.mka import FormTemplate
        row = (
            self.db.query(FormTemplate)
            .filter(FormTemplate.id == template_id, FormTemplate.tenant_id == tenant_id)
            .first()
        )
        if row is None:
            raise LookupError("template not found")
        row.field_mapping = mapping
        self.db.flush()
        return row

    def activate(self, *, tenant_id: UUID, template_id: UUID) -> Any:
        from app.models.mka import FormDefinition, FormTemplate
        from datetime import datetime, timezone

        row = (
            self.db.query(FormTemplate)
            .filter(FormTemplate.id == template_id, FormTemplate.tenant_id == tenant_id)
            .first()
        )
        if row is None:
            raise LookupError("template not found")
        # supersede other active for same form_key
        others = (
            self.db.query(FormTemplate)
            .filter(
                FormTemplate.tenant_id == tenant_id,
                FormTemplate.form_key == row.form_key,
                FormTemplate.status == "active",
                FormTemplate.id != row.id,
            )
            .all()
        )
        for other in others:
            other.status = "superseded"
        row.status = "active"
        row.effective_from = datetime.now(timezone.utc)
        # bind to form definition if present
        definition = (
            self.db.query(FormDefinition)
            .filter(
                FormDefinition.tenant_id == tenant_id,
                FormDefinition.form_key == row.form_key,
                FormDefinition.status == "active",
            )
            .first()
        )
        if definition:
            definition.active_template_id = row.id
        self.db.flush()
        return row

    def preview(
        self,
        *,
        tenant_id: UUID,
        template_id: UUID,
        values: Dict[str, Any],
    ) -> Tuple[bytes, str, str]:
        from app.models.mka import FormTemplate

        row = (
            self.db.query(FormTemplate)
            .filter(FormTemplate.id == template_id, FormTemplate.tenant_id == tenant_id)
            .first()
        )
        if row is None:
            raise LookupError("template not found")
        content = self._storage().get_bytes(row.storage_key)
        mapping = row.field_mapping or {}
        if row.format == "docx":
            out = render_docx_template(content, values, mapping)
            return out, f"{row.form_key}_preview.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if row.format == "xlsx":
            out = render_xlsx_template(content, values, mapping)
            return out, f"{row.form_key}_preview.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        raise ValueError(f"unsupported format: {row.format}")

    def get_active(self, tenant_id: UUID, form_key: str) -> Optional[Any]:
        from app.models.mka import FormTemplate
        return (
            self.db.query(FormTemplate)
            .filter(
                FormTemplate.tenant_id == tenant_id,
                FormTemplate.form_key == form_key,
                FormTemplate.status == "active",
            )
            .order_by(FormTemplate.effective_from.desc().nullslast())
            .first()
        )
