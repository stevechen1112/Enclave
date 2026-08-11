# 產生 T01 報價單 DOCX、T02 異常報告 XLSX、T03 交接班紀錄 XLSX
# 佔位符語法 {{field}} 對齊 app/services/form_template_service.py 的 _PLACEHOLDER_RE
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

OUT = Path(__file__).parent

# ── T01 報價單 DOCX ──────────────────────────────────────────
doc = Document()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("精聯精密工業股份有限公司\n報 價 單 QUOTATION")
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph("客戶名稱：{{customer}}")
doc.add_paragraph("報價日期：{{quote_date}}　　有效期限：{{valid_until}}")
doc.add_paragraph("付款條件：{{payment_terms}}")

table = doc.add_table(rows=2, cols=5)
table.style = "Table Grid"
headers = ["品名／料號", "數量", "單價 (NT$)", "小計 (NT$)", "備註"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
row = table.rows[1].cells
row[0].text = "精密張力控制器 {{part_number}}"
row[1].text = "{{quantity}}"
row[2].text = "{{unit_price}}"
row[3].text = "{{subtotal}}"
row[4].text = "{{notes}}"

doc.add_paragraph("未稅金額：NT$ {{subtotal}}")
doc.add_paragraph("稅額（{{tax_rate}}%）：NT$ {{tax}}")
total_p = doc.add_paragraph()
total_run = total_p.add_run("含稅總計：NT$ {{total}}")
total_run.bold = True
total_run.font.size = Pt(14)

doc.add_paragraph("")
doc.add_paragraph("交期：{{lead_time}}")
doc.add_paragraph("備註：{{remarks}}")
doc.add_paragraph("")
doc.add_paragraph("業務擔當：{{sales_rep}}　　主管核准：＿＿＿＿＿＿　　客戶確認：＿＿＿＿＿＿")

doc.save(OUT / "T01_報價單版型.docx")
print("T01_報價單版型.docx 完成")

# ── 共用 XLSX 樣式 ───────────────────────────────────────────
thin = Side(style="thin")
border = Border(left=thin, right=thin, top=thin, bottom=thin)
head_fill = PatternFill("solid", fgColor="D9E1F2")
head_font = Font(bold=True, size=12)
title_font = Font(bold=True, size=16)


def styled_row(ws, row_idx, values, fill=False, bold=False):
    for col_idx, v in enumerate(values, start=1):
        c = ws.cell(row=row_idx, column=col_idx, value=v)
        c.border = border
        c.alignment = Alignment(vertical="center", wrap_text=True)
        if fill:
            c.fill = head_fill
        if bold:
            c.font = head_font


# ── T02 異常報告 XLSX ────────────────────────────────────────
wb = Workbook()
ws = wb.active
ws.title = "異常報告"
ws.merge_cells("A1:D1")
ws["A1"] = "精聯精密工業 現場異常報告單"
ws["A1"].font = title_font
ws["A1"].alignment = Alignment(horizontal="center")

styled_row(ws, 3, ["設備編號", "{{equipment_id}}", "發生位置／產線", "{{location}}"], fill=False)
ws["A3"].fill = head_fill
ws["C3"].fill = head_fill
styled_row(ws, 4, ["發生時間", "{{occurred_at}}", "回報人", "{{reporter}}"])
ws["A4"].fill = head_fill
ws["C4"].fill = head_fill
styled_row(ws, 5, ["異常類別", "{{category}}", "嚴重程度", "{{severity}}"])
ws["A5"].fill = head_fill
ws["C5"].fill = head_fill
styled_row(ws, 6, ["異常狀況描述", "{{description}}", "", ""])
ws["A6"].fill = head_fill
ws.merge_cells("B6:D6")
styled_row(ws, 7, ["已採取之緊急處置", "{{immediate_action}}", "", ""])
ws["A7"].fill = head_fill
ws.merge_cells("B7:D7")
styled_row(ws, 9, ["班長簽核", "", "設備課簽核", ""])

for col, width in zip("ABCD", [18, 30, 18, 30]):
    ws.column_dimensions[col].width = width

wb.save(OUT / "T02_異常報告版型.xlsx")
print("T02_異常報告版型.xlsx 完成")

# ── T03 交接班紀錄 XLSX ──────────────────────────────────────
wb2 = Workbook()
ws2 = wb2.active
ws2.title = "交接班紀錄"
ws2.merge_cells("A1:D1")
ws2["A1"] = "精聯精密工業 交接班紀錄"
ws2["A1"].font = title_font
ws2["A1"].alignment = Alignment(horizontal="center")

styled_row(ws2, 3, ["班次日期", "{{shift_date}}", "班次", "{{shift}}"])
ws2["A3"].fill = head_fill
ws2["C3"].fill = head_fill
styled_row(ws2, 4, ["產線／區域", "{{line}}", "", ""])
ws2["A4"].fill = head_fill
styled_row(ws2, 5, ["交班人", "{{outgoing}}", "接班人", "{{incoming}}"])
ws2["A5"].fill = head_fill
ws2["C5"].fill = head_fill
styled_row(ws2, 6, ["本班生產狀況", "{{production_summary}}", "", ""])
ws2["A6"].fill = head_fill
ws2.merge_cells("B6:D6")
styled_row(ws2, 7, ["未完成事項／待追蹤", "{{pending_issues}}", "", ""])
ws2["A7"].fill = head_fill
ws2.merge_cells("B7:D7")
styled_row(ws2, 8, ["設備注意事項", "{{equipment_notes}}", "", ""])
ws2["A8"].fill = head_fill
ws2.merge_cells("B8:D8")
styled_row(ws2, 10, ["交班人簽名", "", "接班人簽名", ""])

for col, width in zip("ABCD", [20, 30, 14, 20]):
    ws2.column_dimensions[col].width = width

wb2.save(OUT / "T03_交接班紀錄版型.xlsx")
print("T03_交接班紀錄版型.xlsx 完成")
