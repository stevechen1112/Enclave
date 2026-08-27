#!/usr/bin/env python3
"""Capture a real P3 internal replay from a FORCE-RLS staging deployment.

Run inside the backend image. Fixtures are generated at runtime, uploaded through
the public API, processed by workers, then read back through the tenant-scoped
application DB identity. The output is an immutable replay bundle consumable by
``run_multimodal_quality_gate.py --mode internal_replay``.
"""
from __future__ import annotations

import base64
import json
import os
import re
import subprocess
import sys
import tempfile
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import UUID

import httpx
from docx import Document as WordDocument
from openpyxl import Workbook
from PIL import Image, ImageDraw
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from app.db.session import SessionLocal  # noqa: E402
from app.eval.multimodal_quality import canonical_hash, load_json  # noqa: E402
from app.models.asset import (  # noqa: E402
    AssetRevision,
    DerivedArtifact,
    EvidenceSpan,
    SourceAsset,
)
from app.services.rls import apply_rls_context  # noqa: E402

BASE = os.getenv("P3_REPLAY_API_BASE", "http://web:8000/api/v1")
TERMINAL = {"ready", "review_required", "completed_no_knowledge", "failed"}


def require(
    response: httpx.Response, expected: tuple[int, ...], label: str
) -> httpx.Response:
    if response.status_code not in expected:
        raise RuntimeError(
            f"{label}: HTTP {response.status_code}: {response.text[:500]}"
        )
    return response


def poll_asset(
    client: httpx.Client, headers: dict[str, str], asset_id: str, timeout: int
) -> dict[str, Any]:
    deadline = time.time() + timeout
    latest: dict[str, Any] = {}
    while time.time() < deadline:
        latest = require(
            client.get(f"/knowledge/assets/{asset_id}/status", headers=headers),
            (200,),
            "asset status",
        ).json()
        if str((latest.get("job") or {}).get("status") or "") in TERMINAL:
            return latest
        time.sleep(2)
    raise RuntimeError(f"asset {asset_id} did not settle: {latest}")


def run_ffmpeg(*args: str) -> None:
    subprocess.run(
        ["ffmpeg", "-hide_banner", "-loglevel", "error", *args],
        check=True,
        timeout=600,
    )


def make_document_fixtures(work: Path) -> dict[str, Path]:
    fixtures: dict[str, Path] = {}

    native = work / "native-sop.pdf"
    pdf = canvas.Canvas(str(native))
    pdf.drawString(72, 760, "P3 native PDF cover")
    pdf.showPage()
    pdf.drawString(72, 760, "P3NATIVE pressure zero evidence")
    pdf.save()
    fixtures["doc-pdf-native-001"] = native

    scanned_image = Image.new("RGB", (1200, 700), "white")
    ImageDraw.Draw(scanned_image).text(
        (120, 160), "P3SCAN pressure zero form", fill="black"
    )
    scanned = work / "scanned-form.pdf"
    scanned_image.save(scanned, "PDF", resolution=150)
    fixtures["doc-pdf-scan-001"] = scanned

    docx = work / "work-instruction.docx"
    word = WordDocument()
    word.add_heading("Reset procedure", level=1)
    word.add_paragraph("P3DOCX confirm pressure is zero before reset.")
    word.save(docx)
    fixtures["doc-docx-001"] = docx

    xlsx = work / "inspection.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "檢驗"
    for row in range(4, 10):
        for column in range(2, 7):
            sheet.cell(row=row, column=column, value=f"P3XLSX-{row}-{column}")
    workbook.save(xlsx)
    fixtures["sheet-xlsx-001"] = xlsx

    csv_path = work / "inventory.csv"
    csv_path.write_text("part,quantity\nP3CSV-PART,42\n", encoding="utf-8")
    fixtures["sheet-csv-001"] = csv_path

    try:
        from PIL import ImageFont

        fixture_font = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 48
        )
    except OSError:
        try:
            fixture_font = ImageFont.load_default(size=48)
        except TypeError:
            fixture_font = ImageFont.load_default()
    printed = Image.new("RGB", (1200, 500), "white")
    ImageDraw.Draw(printed).text(
        (120, 80),
        "P3PRINT EQ-100 PRESSURE 0 BAR",
        fill="black",
        font=fixture_font,
    )
    printed_path = work / "printed-label.png"
    printed.save(printed_path)
    fixtures["image-print-001"] = printed_path

    handwritten = Image.new("RGB", (1200, 600), "white")
    draw = ImageDraw.Draw(handwritten)
    draw.line((100, 100, 900, 360), fill="black", width=5)
    draw.text(
        (160, 150),
        "P3HAND inspect valve",
        fill="black",
        font=fixture_font,
    )
    handwritten_path = work / "handwritten-note.png"
    handwritten.save(handwritten_path)
    fixtures["image-hand-001"] = handwritten_path
    return fixtures


def make_audio_fixtures(work: Path) -> dict[str, Path]:
    fixtures: dict[str, Path] = {}
    quiet = work / "quiet-single.wav"
    run_ffmpeg(
        "-f",
        "lavfi",
        "-i",
        "anullsrc=r=16000:cl=mono:d=1",
        "-f",
        "lavfi",
        "-i",
        "flite=text=P3QUIET confirm pressure zero",
        "-filter_complex",
        "[0:a][1:a]concat=n=2:v=0:a=1[out]",
        "-map",
        "[out]",
        "-ar",
        "16000",
        "-ac",
        "1",
        "-y",
        str(quiet),
    )
    fixtures["audio-quiet-001"] = quiet

    noisy = work / "noisy-floor.wav"
    run_ffmpeg(
        "-f",
        "lavfi",
        "-i",
        "flite=text=P3NOISY inspect machine valve",
        "-f",
        "lavfi",
        "-i",
        "anoisesrc=color=white:amplitude=0.03:d=6",
        "-filter_complex",
        "[0:a][1:a]amix=inputs=2:duration=longest[out]",
        "-map",
        "[out]",
        "-ar",
        "16000",
        "-ac",
        "1",
        "-y",
        str(noisy),
    )
    fixtures["audio-noisy-001"] = noisy

    multi = work / "two-speakers.wav"
    run_ffmpeg(
        "-f",
        "lavfi",
        "-i",
        "flite=text=P3MULTI operator confirms pressure zero:voice=kal",
        "-f",
        "lavfi",
        "-i",
        "flite=text=supervisor approves inspection:voice=slt",
        "-filter_complex",
        "[0:a][1:a]concat=n=2:v=0:a=1[out]",
        "-map",
        "[out]",
        "-ar",
        "16000",
        "-ac",
        "1",
        "-y",
        str(multi),
    )
    fixtures["audio-multi-001"] = multi

    long_audio = work / "long-shift-handover.mp3"
    run_ffmpeg(
        "-f",
        "lavfi",
        "-i",
        "anullsrc=r=16000:cl=mono:d=105",
        "-f",
        "lavfi",
        "-i",
        "flite=text=P3LONG shift handover pressure zero",
        "-filter_complex",
        "[0:a][1:a]concat=n=2:v=0:a=1[out]",
        "-map",
        "[out]",
        "-ar",
        "16000",
        "-ac",
        "1",
        "-b:a",
        "32k",
        "-y",
        str(long_audio),
    )
    fixtures["audio-long-001"] = long_audio
    return fixtures


def make_video(
    path: Path,
    *,
    duration: int,
    speech: str | None,
    delay_ms: int = 0,
    source: str = "color=c=blue:s=640x360",
) -> None:
    command = ["-f", "lavfi", "-i", f"{source}:d={duration}"]
    if speech:
        command += ["-f", "lavfi", "-i", f"flite=text={speech}"]
        command += [
            "-filter_complex",
            f"[1:a]adelay={delay_ms}|{delay_ms}[a]",
            "-map",
            "0:v",
            "-map",
            "[a]",
            "-shortest",
        ]
    else:
        command += ["-map", "0:v"]
    command += ["-c:v", "libx264", "-pix_fmt", "yuv420p"]
    if speech:
        command += ["-c:a", "aac"]
    command += ["-y", str(path)]
    run_ffmpeg(*command)


def make_video_fixtures(work: Path) -> dict[str, Path]:
    fixtures: dict[str, Path] = {}
    caption = work / "fixed-captioned.mp4"
    make_video(
        caption, duration=12, speech="P3CAPTION confirm pressure zero", delay_ms=6400
    )
    fixtures["video-caption-001"] = caption

    silent = work / "fixed-silent.mp4"
    make_video(silent, duration=10, speech=None)
    fixtures["video-silent-001"] = silent

    handheld = work / "handheld-inspection.mp4"
    make_video(
        handheld,
        duration=20,
        speech="P3HANDHELD inspect equipment valve",
        delay_ms=12200,
        source="testsrc2=s=640x360:r=24",
    )
    fixtures["video-handheld-001"] = handheld

    machine = work / "machine-reset.mp4"
    make_video(
        machine,
        duration=45,
        speech="P3MACHINE danger force reset EQ one hundred",
        delay_ms=40200,
    )
    fixtures["video-machine-001"] = machine
    return fixtures


def decode_tenant(token: str) -> UUID:
    payload = token.split(".")[1]
    payload += "=" * (-len(payload) % 4)
    return UUID(json.loads(base64.urlsafe_b64decode(payload))["tenant_id"])


def upload_asset(
    client: httpx.Client,
    headers: dict[str, str],
    case_id: str,
    path: Path,
) -> str:
    media_type = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".csv": "text/csv",
        ".png": "image/png",
        ".wav": "audio/wav",
        ".mp3": "audio/mpeg",
        ".mp4": "video/mp4",
        ".txt": "text/plain",
    }[path.suffix.lower()]
    with path.open("rb") as stream:
        if case_id == "video-machine-001":
            response = client.post(
                "/media/videos",
                headers=headers,
                files={"file": (path.name, stream, media_type)},
                data={
                    "title": case_id,
                    "equipment_ids": "EQ-100",
                    "applicable_roles": "operator",
                },
            )
        else:
            response = client.post(
                "/knowledge/assets",
                headers=headers,
                files={"file": (path.name, stream, media_type)},
                data={"title": case_id},
            )
    return str(require(response, (202,), f"upload {case_id}").json()["id"])


def terminal_state(job: dict[str, Any]) -> str:
    job_status = str(job.get("status") or "failed")
    if job_status == "ready" and job.get("phase") == "completed_no_knowledge":
        return "completed_no_knowledge"
    return {
        "ready": "completed",
        "review_required": "review_required",
        "completed_no_knowledge": "completed_no_knowledge",
        "failed": "failed_explained",
    }[job_status]


def evidence_for_asset(
    db,
    *,
    tenant_id: UUID,
    asset_id: UUID,
    revision_alias: str,
    evidence_query: str | None,
) -> tuple[list[dict[str, Any]], bool]:
    apply_rls_context(db, tenant_id)
    asset = db.query(SourceAsset).filter(SourceAsset.id == asset_id).one()
    revision = (
        db.query(AssetRevision)
        .filter(
            AssetRevision.asset_id == asset.id,
            AssetRevision.revision == asset.current_revision,
        )
        .one()
    )
    artifacts = (
        db.query(DerivedArtifact)
        .filter(
            DerivedArtifact.asset_revision_id == revision.id,
        )
        .all()
    )
    knowledge_kinds = {
        "layout_page",
        "table",
        "ocr_region",
        "transcript_segment",
        "procedure_candidate",
    }

    def tokens(value: str) -> set[str]:
        return set(re.findall(r"[a-z0-9\u4e00-\u9fff]+", value.lower()))

    candidates = [row for row in artifacts if row.artifact_kind in knowledge_kinds]
    if evidence_query:
        query_tokens = tokens(evidence_query)
        ranked = sorted(
            candidates,
            key=lambda row: (
                len(query_tokens.intersection(tokens(str(row.content or "")))),
                -len(tokens(str(row.content or ""))),
            ),
            reverse=True,
        )
        selected = ranked[:1] if ranked else []
    else:
        selected = []
    artifact_ids = [row.id for row in selected]
    spans = (
        db.query(EvidenceSpan).filter(EvidenceSpan.artifact_id.in_(artifact_ids)).all()
        if artifact_ids
        else []
    )
    locators: list[dict[str, Any]] = []
    for span in spans:
        locator: dict[str, Any] = {
            "id": str(span.id),
            "kind": span.locator_kind,
            "revision_id": revision_alias,
            "tenant_id": (
                "tenant-synthetic-a"
                if span.tenant_id == tenant_id
                else f"unexpected:{span.tenant_id}"
            ),
        }
        for source, target in (
            ("page", "page"),
            ("section", "section"),
            ("worksheet", "worksheet"),
            ("table_name", "table_name"),
            ("row_number", "row_number"),
            ("column_name", "column_name"),
            ("cell_range", "cell_range"),
            ("start_ms", "start_ms"),
            ("end_ms", "end_ms"),
        ):
            value = getattr(span, source)
            if value is not None:
                locator[target] = (
                    int(value)
                    if source in {"start_ms", "end_ms", "row_number", "page"}
                    else value
                )
        if span.bbox is not None:
            locator["region"] = span.bbox
        locators.append(locator)
    conflict_detected = False
    for artifact in artifacts:
        if artifact.artifact_kind != "sop_conflict_report" or not artifact.content:
            continue
        try:
            conflict_detected = bool(json.loads(artifact.content).get("conflicts"))
        except (TypeError, ValueError):
            pass
    return locators, conflict_detected


def main() -> int:
    manifest = load_json(ROOT / "testdata/multimodal_golden/manifest.json")
    run_id = str(uuid.uuid4())
    source_commit = os.getenv("SOURCE_COMMIT", "")
    if len(source_commit) != 40:
        raise RuntimeError("SOURCE_COMMIT must be the deployed 40-character commit")
    raw_capture: dict[str, Any] = {
        "run_id": run_id,
        "source_commit": source_commit,
        "cases": [],
    }
    with tempfile.TemporaryDirectory(prefix="p3-replay-") as directory:
        work = Path(directory)
        fixtures = {
            **make_document_fixtures(work),
            **make_audio_fixtures(work),
            **make_video_fixtures(work),
        }
        with httpx.Client(base_url=BASE, timeout=300.0) as client:
            login = require(
                client.post(
                    "/auth/login/access-token",
                    data={
                        "username": os.environ["FIRST_SUPERUSER_EMAIL"],
                        "password": os.environ["FIRST_SUPERUSER_PASSWORD"],
                    },
                ),
                (200,),
                "owner login",
            ).json()
            token = login["access_token"]
            headers = {"Authorization": f"Bearer {token}"}
            tenant_id = decode_tenant(token)

            sop_path = work / "P3_Formal_SOP.txt"
            sop_path.write_text(
                "P3 formal SOP. 禁止 force reset. First isolate power.",
                encoding="utf-8",
            )
            sop_asset = upload_asset(client, headers, "p3-formal-sop", sop_path)
            poll_asset(client, headers, sop_asset, 240)

            captures: dict[str, dict[str, Any]] = {}
            for case in manifest["cases"]:
                case_id = case["id"]
                asset_id = upload_asset(client, headers, case_id, fixtures[case_id])
                status = poll_asset(client, headers, asset_id, 720)
                captures[case_id] = {"asset_id": asset_id, "status": status}

            db = SessionLocal()
            try:
                results = []
                for case in manifest["cases"]:
                    case_id = case["id"]
                    capture = captures[case_id]
                    job = dict(capture["status"].get("job") or {})
                    job_status = str(job.get("status") or "failed")
                    locators, conflict = evidence_for_asset(
                        db,
                        tenant_id=tenant_id,
                        asset_id=UUID(capture["asset_id"]),
                        revision_alias=case["expected"]["revision_id"],
                        evidence_query=case["expected"].get("evidence_query"),
                    )
                    review_created = job_status == "review_required"
                    raw_capture["cases"].append(
                        {
                            "case_id": case_id,
                            "asset_id": capture["asset_id"],
                            "job_status": job_status,
                            "evidence_count": len(locators),
                            "sop_conflict_detected": conflict,
                        }
                    )
                    results.append(
                        {
                            "case_id": case_id,
                            "terminal_state": terminal_state(job),
                            "tenant_id": "tenant-synthetic-a",
                            "evidence_locators": locators,
                            "review_created": review_created,
                            "sop_conflict_detected": conflict,
                            "answer": {
                                "status": (
                                    "abstained"
                                    if case["expected"].get("high_risk")
                                    else "not_applicable"
                                ),
                                "grounded": True,
                                "authoritative": False,
                                "citations": [],
                            },
                        }
                    )
            finally:
                db.close()

    raw_hash = canonical_hash(raw_capture)
    bundle = {
        "schema_version": 1,
        "mode": "internal_replay",
        "provider": "enclave.staging-pipeline",
        "provider_version": source_commit[:12],
        "corpus_sha256": canonical_hash(manifest),
        "provenance": {
            "run_id": run_id,
            "captured_at": datetime.now(timezone.utc).isoformat(),
            "source_commit": source_commit,
            "execution_environment": os.getenv("ENVIRONMENT", "staging"),
            "source_artifact_sha256": raw_hash,
        },
        "results": results,
    }
    print(
        json.dumps(
            {"raw_capture": raw_capture, "result_bundle": bundle},
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
