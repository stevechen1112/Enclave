"""從原始檔案抽取文字，建立盲測 ground truth（獨立於系統解析結果）。"""
import io
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

FILES = {
    "tax113": r"C:\Users\User\Desktop\八策\113年營所稅申報書_E42八策.pdf",
    "quote_lafa": r"C:\Users\User\Desktop\八策\1140213-拉法(存奕官網)_報價單.docx",
    "resume_wu": r"C:\Users\User\Downloads\吳文曄-履歷.pdf",
    "giant_report": r"C:\Users\User\Documents\巨大機械9921深度研究報告.pdf",
    "amazon_quote": r"C:\Users\User\Desktop\客戶\光昱金屬\亞馬遜行銷報價.pdf",
    "taipei_subsidy": r"C:\Users\User\Desktop\八策\政府補助\臺北市產業發展獎勵補助計畫申請簡介.pdf",
    "manual": r"C:\Users\User\Downloads\基礎操作教學手冊.pdf",
    "xunyun": r"C:\Users\User\Desktop\客戶\巽耘\【將能數位行銷】-巽耘法律事務所-健診報告.pdf",
}

for key, path in FILES.items():
    out = f"artifacts/_blind_src_{key}.txt"
    try:
        if path.lower().endswith(".docx"):
            import docx
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(parts)
        else:
            import pdfplumber
            parts = []
            with pdfplumber.open(path) as pdf:
                for pg in pdf.pages[:15]:
                    parts.append(pg.extract_text() or "")
            text = "\n".join(parts)
        with open(out, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"{key}: {len(text)} chars -> {out}")
    except Exception as e:
        print(f"{key}: ERROR {e}")
